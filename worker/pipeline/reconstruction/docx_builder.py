import io

from docx import Document as DocxDocument
from docx.shared import Pt


def build_docx(chunks: list[tuple[str, str]], filename: str) -> bytes:
    """Build a DOCX document from translated chunks.

    Args:
        chunks: List of (original_text, translated_text) pairs.
        filename: Original filename (used in the document title).

    Returns:
        DOCX as bytes.
    """
    doc = DocxDocument()

    heading = doc.add_heading(f"TRANSLATION \u2014 {filename}", level=1)
    if heading.runs:
        heading.runs[0].font.size = Pt(14)

    for _, translated in chunks:
        for para_text in translated.split("\n\n"):
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                if p.runs:
                    p.runs[0].font.size = Pt(11)
        doc.add_paragraph()  # blank line between chunks

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
