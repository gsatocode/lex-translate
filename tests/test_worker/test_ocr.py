import io

import pytest
from docx import Document as DocxDocument
from reportlab.pdfgen.canvas import Canvas
from unittest.mock import patch

from worker.pipeline.ocr.base import OCRAdapter, OCRResult
from worker.pipeline.ocr.pdf_adapter import PDFOCRAdapter
from worker.pipeline.ocr.pdfplumber_adapter import PDFPlumberAdapter
from worker.pipeline.ocr.docx_adapter import DocxAdapter
from worker.pipeline.ocr.paddle_adapter import PaddleOCRAdapter


@pytest.mark.asyncio
async def test_pdfplumber_extracts_text():
    buf = io.BytesIO()
    c = Canvas(buf)
    c.drawString(72, 720, "The applicant was born on 01/01/1990.")
    c.save()
    pdf_bytes = buf.getvalue()

    adapter = PDFPlumberAdapter()
    result = await adapter.extract(pdf_bytes)

    assert isinstance(result, OCRResult)
    assert "01/01/1990" in result.text
    assert result.page_count == 1


@pytest.mark.asyncio
async def test_pdfplumber_empty_pdf_returns_empty_text():
    buf = io.BytesIO()
    c = Canvas(buf)
    c.showPage()
    c.save()
    pdf_bytes = buf.getvalue()

    adapter = PDFPlumberAdapter()
    result = await adapter.extract(pdf_bytes)

    assert result.text == ""
    assert result.page_count == 1


@pytest.mark.asyncio
async def test_docx_adapter_extracts_paragraphs():
    doc = DocxDocument()
    doc.add_paragraph("First paragraph of legal document.")
    doc.add_paragraph("Second paragraph with date 15/03/2024.")
    buf = io.BytesIO()
    doc.save(buf)

    adapter = DocxAdapter()
    result = await adapter.extract(buf.getvalue())

    assert "First paragraph" in result.text
    assert "15/03/2024" in result.text
    assert result.page_count == 1


@pytest.mark.asyncio
async def test_docx_adapter_skips_empty_paragraphs():
    doc = DocxDocument()
    doc.add_paragraph("")
    doc.add_paragraph("Only real content.")
    doc.add_paragraph("")
    buf = io.BytesIO()
    doc.save(buf)

    adapter = DocxAdapter()
    result = await adapter.extract(buf.getvalue())

    assert result.text == "Only real content."


@pytest.mark.asyncio
async def test_paddle_adapter_raises_without_package():
    adapter = PaddleOCRAdapter()
    with patch.dict("sys.modules", {"paddleocr": None}):
        with pytest.raises(RuntimeError, match="PaddleOCR not installed"):
            await adapter.extract(b"fake image bytes")


class _StubAdapter(OCRAdapter):
    def __init__(self, result: OCRResult):
        self.result = result
        self.calls = 0

    async def extract(self, data: bytes) -> OCRResult:
        self.calls += 1
        return self.result


@pytest.mark.asyncio
async def test_pdf_adapter_prefers_embedded_text():
    text_adapter = _StubAdapter(OCRResult(text="Certified translation of the attached document.", page_count=1))
    scanned_adapter = _StubAdapter(OCRResult(text="Fallback OCR", page_count=1))

    adapter = PDFOCRAdapter(text_adapter=text_adapter, scanned_adapter=scanned_adapter, min_text_chars=20)
    result = await adapter.extract(b"%PDF-1.4")

    assert result.text == "Certified translation of the attached document."
    assert text_adapter.calls == 1
    assert scanned_adapter.calls == 0


@pytest.mark.asyncio
async def test_pdf_adapter_falls_back_to_ocr_when_pdf_text_is_empty():
    text_adapter = _StubAdapter(OCRResult(text="", page_count=1))
    scanned_adapter = _StubAdapter(OCRResult(text="OCR output from scanned PDF", page_count=1))

    adapter = PDFOCRAdapter(text_adapter=text_adapter, scanned_adapter=scanned_adapter, min_text_chars=20)
    result = await adapter.extract(b"%PDF-1.4")

    assert result.text == "OCR output from scanned PDF"
    assert text_adapter.calls == 1
    assert scanned_adapter.calls == 1


@pytest.mark.asyncio
async def test_pdf_adapter_falls_back_to_ocr_when_pdf_text_layer_is_low_signal():
    text_adapter = _StubAdapter(OCRResult(text="I\nI\nI", page_count=1))
    scanned_adapter = _StubAdapter(OCRResult(text="Recovered full certificate text", page_count=1))

    adapter = PDFOCRAdapter(text_adapter=text_adapter, scanned_adapter=scanned_adapter, min_text_chars=20)
    result = await adapter.extract(b"%PDF-1.4")

    assert result.text == "Recovered full certificate text"
    assert scanned_adapter.calls == 1
